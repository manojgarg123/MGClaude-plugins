#!/usr/bin/env python3
"""
extract_text.py — Extract plain text from PDF or DOCX files.
Usage: python extract_text.py "<file_path>"
Prints extracted text to stdout.
"""

import sys
import os

def extract_from_pdf(path):
    try:
        import pdfplumber
        text = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text.append(t)
        return "\n".join(text)
    except ImportError:
        # Fallback to pypdf
        try:
            from pypdf import PdfReader
            reader = PdfReader(path)
            text = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    text.append(t)
            return "\n".join(text)
        except ImportError:
            raise RuntimeError("Please install pdfplumber: pip install pdfplumber --break-system-packages")

def extract_from_docx(path):
    try:
        from docx import Document
        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_text.append(row_text)
        return "\n".join(paragraphs + table_text)
    except ImportError:
        raise RuntimeError("Please install python-docx: pip install python-docx --break-system-packages")

def extract_text(path):
    if not os.path.exists(path):
        raise FileNotFoundError(f"File not found: {path}")

    ext = os.path.splitext(path)[1].lower()

    if ext == ".pdf":
        return extract_from_pdf(path)
    elif ext in (".docx", ".doc"):
        return extract_from_docx(path)
    elif ext == ".txt":
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported types: PDF, DOCX, TXT")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python extract_text.py <file_path>", file=sys.stderr)
        sys.exit(1)

    file_path = sys.argv[1]
    try:
        text = extract_text(file_path)
        print(text)
    except Exception as e:
        print(f"ERROR: {e}", file=sys.stderr)
        sys.exit(1)
